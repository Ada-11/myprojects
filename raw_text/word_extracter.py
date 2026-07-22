import sys
import os
import docx

def save_docx_extract(docx_path, output_txt_path="/Users/ada/myprojects/my-first-app/raw_text/claims_process.txt"):
    """Extracts text paragraphs and structured tables from a Word file."""
    if not os.path.exists(docx_path):
        print(f"[ERROR] The file '{docx_path}' could not be found.")
        return

    full_text = []
    print(f"[PROCESSING] Word document: {docx_path}...")
    doc = docx.Document(docx_path)

    # 1. Extract standard paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)

    # 2. Extract table grids
    for table in doc.tables:
        full_text.append("\n--- Table Data ---")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            cleaned_row = []
            for text in row_text:
                if not cleaned_row or text != cleaned_row[-1]:
                    cleaned_row.append(text)
            full_text.append(" | ".join(cleaned_row))

    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))
    print(f"[SUCCESS] Extraction saved to: {output_txt_path}")

if __name__ == "__main__":
    target_file = sys.argv[1] if len(sys.argv) > 1 else "/Users/ada/myprojects/my-first-app/raw_text/claims-process.docx"
    save_docx_extract(target_file)